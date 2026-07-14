from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = 'Grocery_Change_Log.docx'

def font(run, size=10.5, bold=False, color=None):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def heading(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), 13, True, (46,116,181))

def body(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.05; font(p.add_run(text),10)

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run('PROJECT HUB  |  CHANGE LOG'),8.5,True,(89,89,89))
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run('GroceryLedger  •  Maintained in Asia/Kolkata'),8,False,(89,89,89))
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2); font(p.add_run('GroceryLedger Change Log'),20,True,(11,37,69))
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(7); font(p.add_run('Dated record of product decisions, implementation changes, verification, and known limits.'),11,False,(85,85,85))

heading(doc,'12 July 2026, 01:30 IST — Receipt-driven grocery ledger foundation')
body(doc,'What changed: Created the native SwiftUI/SwiftData iPhone project; added local PDF storage, equal Ekta/Ritesh splitting, personal-item exclusion, balances, manual settlements, invoice review, duplicate-invoice protection, and simulator invoice staging.')
body(doc,'Why: Establish a working personal tool before expanding scope. Receipt data remains editable because parsing is not trusted as the final record.')
body(doc,'Verification: Simulator build succeeded repeatedly. Instamart coordinate-based extraction was smoke-tested against a supplied invoice. Known limit: Blinkit product labels may be abbreviated and invoice-date extraction requires continued refinement.')

heading(doc,'12 July 2026, 01:55 IST — Suggestion logic and data-quality findings')
body(doc,'What changed: Added repeat-product suggestions and investigated the imported simulator data. Saved history contains repeated items including Paneer, Desi Tomato, Green Chilli, Banana, and Potato.')
body(doc,'Why: The original due-only presentation hid useful repeat evidence. The revised direction is to show repeated products as possible buys and label the evidence, never claim that an item is certainly finished.')
body(doc,'Known limit: The first cadence rule used an average interval and is not suitable for sparse history. Confirmed replacement: use the most recent completed interval for every item; with a longer history, present a robust median/range instead of an average.')

heading(doc,'12 July 2026, 02:00 IST — Confirmed next scope')
body(doc,'Cadence: Implement the last-two-purchase interval rule for every product. Expiry: add an explicitly estimated use-by date when no label date exists; never present it as a known expiry. Food: supplied Zomato order summary (Order 8345250035) has restaurant, item lines, charges, discount, total, buyer, and order time, so it is suitable for a separate food-order parser and a Splitwise-style shared-expense review. Manual entry: add a general expense form for groceries, food, Wi-Fi, water, and other roommate costs.')
body(doc,'Documentation policy: Grocery_Change_Log.docx and Grocery_Slip_Tracker_Market_Research.docx must be updated for material changes with an Asia/Kolkata timestamp, rationale, verification outcome, and known limitation. A scheduled documentation check runs daily at 09:00 and 19:00 IST.')

heading(doc,'12 July 2026, 02:10 IST — Data-minimization privacy rule')
body(doc,'Decision: Do not persist raw receipts, customer names, addresses, payment method, payment instrument details, payment references, or payment status. The importer reads a chosen PDF only in memory to extract the minimum ledger fields, then discards it.')
body(doc,'Retained data: merchant or bill label, category, date, total, line-item label/amount when needed for a split or suggestion, who in the local group paid, and settlement amount/date. The group payer is not a payment method or account detail.')
body(doc,'Follow-up required: existing simulator receipts and legacy raw-PDF blobs may contain personal data and require an explicit cleanup action before the prototype is considered compliant with this rule.')

heading(doc,'12 July 2026, 02:55 IST — Cadence calculation correction')
body(doc,'Correction: replaced the still-active average-gap calculation with the gap between the latest two distinct purchase dates for each item. Why: an average could show Paneer’s 13-day interval despite its most recent buying cycle being different. Verification: iPhone simulator build succeeded after the correction.')

heading(doc,'12 July 2026, 13:20 IST — Parser and suggestion-data repair')
body(doc,'Correction: the older Blinkit parser kept only trailing words, producing labels such as “Powder(Pack)”. The repaired parser extracts full product labels from the receipt table. The suggestion engine now collapses repeated lines per invoice and withholds any cadence when duplicate same-day imports make the purchase history unreliable. Verification: tested against all six supplied Blinkit receipts and installed on the existing simulator; the false 13-day cadence no longer appears. Limit: existing malformed saved labels remain until records are cleaned up or re-imported.')

heading(doc,'12 July 2026, 13:29 IST — Authorized simulator data cleanup')
body(doc,'Action: removed the simulator app and reinstalled the corrected build, deleting all prior purchases, balances, malformed item labels, legacy raw-PDF blobs, and staged receipt files. Why: the imported history was both unreliable for forecasting and inconsistent with the data-minimization rule. Verification: the fresh app opens “All settled up” with no purchases; its Documents folder contains no PDF files. Original source PDFs in Downloads were not changed.')

heading(doc,'12 July 2026, 13:35 IST — Temporary simulator import staging')
body(doc,'Action: copied the ten supplied grocery PDFs from Downloads into a temporary simulator Files folder so the clean app can import them. Privacy guardrail: this is staging only, not app persistence; the importer discards each source PDF after parsing. Follow-up: delete the temporary Files folder after the imports are reviewed.')

heading(doc,'12 July 2026, 13:45 IST — Files visibility configuration repaired')
body(doc,'Root cause: the generated app manifest omitted UIFileSharingEnabled even though the project setting stated it. Repair: added an explicit Info.plist with UIFileSharingEnabled, LSSupportsOpeningDocumentsInPlace, and the required CFBundleExecutable. Verification: the simulator accepted the rebuilt app, and its installed manifest contains both Files-sharing keys. The temporary receipts are exposed under On My iPhone > Grocery Ledger > Import PDFs.')

heading(doc,'12 July 2026, 18:25 IST — Shared-expense and food-order expansion')
body(doc,'What changed: added manual purchase entry for Groceries, Food, Wi-Fi, Water, Household, and Other categories. A manually added grocery or household purchase can carry an optional estimated use-by date. Purchase records now display category and, where supplied, the estimated date. Why: a shared household has recurring costs that do not originate from a grocery invoice, and invoice expiry is usually absent.')
body(doc,'Food parser: added deterministic support for the supplied Zomato food-order summary. It identifies La Pino\'z Pizza as the merchant, Food as the category, the order time, five line items, and the ₹674.88 final total. Verification: compiled a local parser harness against Order_ID_8345250035.pdf, then ran an iPhone Simulator SDK build; both succeeded.')
body(doc,'Limit: the estimated use-by date is a user-entered estimate, not a known package expiry. The app does not yet look up barcode/product data, read package labels, schedule local notifications, or infer an expiry from a receipt. Imported invoices still require the existing review flow before save.')

heading(doc,'12 July 2026, 19:08 IST — Verified SwiftData migration repair')
body(doc,'Root cause: eight receipt purchases had been saved in the simulator database, but the app showed “No purchases” after a new required category field was introduced. SwiftData reported: “Cannot migrate store in-place: Validation error missing attribute values on mandatory destination attribute.” The failure was a data-schema migration, not PDF selection or receipt saving.')
body(doc,'Repair: added the safe default “Groceries” for the newly introduced Purchase.category field, then rebuilt and launched the simulator app. Verification: the migrated SQLite schema contains ZCATEGORY; all eight existing purchases have category Groceries; the live simulator dashboard displays the restored balance and suggestions; and an isolated SwiftData insert/save/fetch check passed. The importer also performs an explicit save and reports any future save error.')

heading(doc,'12 July 2026, 20:43 IST — Persistence verification harness recorded')
body(doc,'What changed: added a dedicated Tools/VerifyPersistence.swift harness that creates an in-memory SwiftData container, inserts a purchase plus item, saves it, fetches it back, and asserts that the merchant, category, and item relationship survive persistence.')
body(doc,'Why: the migration repair needed a repeatable verification artifact instead of a one-off manual claim. This isolates the ledger model from UI/import flows and gives the project a concrete persistence smoke test for future schema changes.')
body(doc,'Verification: source inspection confirms the harness covers save and fetch behavior against Purchase, PurchaseItem, and Settlement models with the new category default in place. Known limit: this run could not execute SwiftData macro builds from the sandbox, so the harness still needs normal local execution in Xcode or `swift run` outside the sandbox.')

heading(doc,'13 July 2026, 11:56 IST — Verified notification preferences, restock scope, and regression suite')
body(doc,'What changed: added an iPhone Settings tab with Off, Daily, and Weekly reminder choices plus a user-selected time and optional weekday. Reminders use a neutral “Review possible buys” message; they do not assert that anything is definitely finished. The restock engine now considers only non-personal grocery/household items explicitly marked for restock, excludes Food, and gives a manually entered estimated use-by date precedence over an interval prediction.')
body(doc,'Why: notification frequency must be a user decision, food should never pollute grocery restock predictions, and a known user-entered estimate is more direct than a derived buying interval. The app preserves the existing policy that use-by values are estimates rather than verified package expiry dates.')
body(doc,'Verification: an iPhone Simulator SDK build succeeded. Added GroceryLedgerTests and executed 7 simulator tests successfully: shared/personal split math, no-debt-after-deletion safeguard, latest-two-date interval, estimated-use-by precedence, Food/untracked exclusion, food parser fields using sanitised text, and raw-PDF non-persistence by default.')

heading(doc,'13 July 2026, 11:56 IST — Browser testing version created')
body(doc,'What changed: created the web/Grocery Ledger browser implementation for eventual GitHub Pages sharing. It provides browser-local manual expenses, categories for groceries/food/Wi-Fi/water/household/other, personal-item exclusion, balance settlement records, a transparent possible-buy calculation, demo data, and a clear-data action. It intentionally accepts no receipt uploads in this testing version.')
body(doc,'Privacy: the web version uses only the browser’s local storage for its reviewed ledger. It does not upload or save PDFs, addresses, payment mode, card/UPI details, payment references, or payment status. It includes a visible privacy notice and a local-only test-data boundary.')
body(doc,'Verification: generated a project social-preview card at web/public/og.png; browser build completed successfully; its local privacy/control test passed. Known limit: GitHub Pages cannot be published until the project root is connected to the GitHub repository rather than the accidental empty nested repository at 359/359.')

heading(doc,'13 July 2026, 12:03 IST — GitHub Pages deployment package prepared')
body(doc,'What changed: added a standalone static website package in docs/ plus .github/workflows/pages.yml. The GitHub Pages version has the same browser-local manual ledger, balance, settlement, personal-item, and restock-cue behavior; it does not depend on a server or receipt upload. The workflow publishes docs/ when main is pushed.')
body(doc,'Why: GitHub Pages serves static files, whereas the richer web/ prototype is built for a worker runtime. A separate static package ensures a shared test URL can work without moving any personal receipt, address, or payment data onto a public host.')
body(doc,'Verification: JavaScript syntax check passed; the static asset, GitHub Pages workflow, and social-preview image are present. Publishing is deliberately blocked until the real project root is initialised as the Git repository and connected to github.com/EktaDhan/359; the existing nested 359/359 repository is empty and must not be used as the project root.')
heading(doc,'13 July 2026, 12:54 IST — Public browser tester deployed')
body(doc,'What changed: the real project repository was connected to GitHub using SSH, its local and remote histories were merged, and the static docs/ package was published from the main branch through GitHub Pages.')
body(doc,'Verification: the live Grocery Ledger tester rendered successfully at https://ektadhan.github.io/359/. The rendered page shows the browser-only privacy promise, manual expense and settlement controls, a balance summary, and a restock cue. No receipt PDFs, addresses, payment modes, card information, or UPI details are part of the deployed static site.')
heading(doc,'13 July 2026, 13:12 IST — Browser settlement validation repaired')
body(doc,'What changed: corrected the static web settlement form so the hidden required purchase-label field is no longer validated while recording a payment. Previously, browser validation silently blocked Save for a settlement because that field was irrelevant but still required.')
body(doc,'Verification: on a clean local static preview, loaded demo data and recorded Ekta paying Rs 300.00. The dialog closed, the confirmation read “Settlement recorded on this browser only,” and the balance changed from Ekta owing Ritesh Rs 367.00 to Rs 67.00. This correction must be committed and pushed before it reaches the public tester.')
heading(doc,'13 July 2026, 13:18 IST — Browser dialog Close and Cancel repaired')
body(doc,'What changed: converted the settlement dialog Close and Cancel controls to explicit non-submitting buttons. Before this repair, native form validation also blocked those controls whenever the hidden purchase label was required.')
body(doc,'Verification: opened the settlement dialog in a local static preview, tested Close, reopened it, then tested Cancel. Both actions dismissed the dialog with zero dialog instances remaining. The prior Save verification remains valid. This second correction must be included in the same follow-up commit and push.')
heading(doc,'13 July 2026, 13:40 IST — Shared ledger architecture approved')
body(doc,'Decision: the website will become a two-person shared ledger for Ekta and Ritesh, rather than remain a single-device browser-only tool. Supabase Free was selected as the backend, with passwordless email magic links or one-time codes for separate sign-in.')
body(doc,'Privacy boundary: the backend may hold account identity, household membership, reviewed ledger entries, settlement records, and restock settings. It must not hold raw receipt PDFs, addresses, payment mode, card/UPI details, bank details, or payment references. Row Level Security must ensure that only household members can access that household data.')
body(doc,'Next dependency: Ekta must create the Supabase project and supply only its project URL and publishable/anon key. No service-role key, secret, password, or email code should be shared with Codex.')
heading(doc,'13 July 2026, 13:42 IST — Secure shared-sync schema prepared')
body(doc,'What changed: added supabase/schema.sql for the approved shared-ledger design. It creates households, household membership, reviewed purchases, and settlements; supports one-time-code/magic-link authenticated users; provides household create/join RPCs; enables Row Level Security; and registers the shared tables for real-time updates.')
body(doc,'Security verification: the schema contains no columns for receipts, address, payment mode, payment reference, bank/card/UPI detail, or database secrets. Access is scoped to household membership. The project owner ran it successfully in the Supabase SQL Editor; the public website remains disconnected until authentication redirect settings and sign-in verification are complete.')
heading(doc,'13 July 2026, 14:08 IST — Shared web sign-in and sync client added')
body(doc,'What changed: configured the static website with the Supabase project URL and publishable key, added passwordless email sign-in, household create/join via invite code, household-scoped purchases and settlements, payment history, and real-time ledger reloads. The UI now makes it clear that each person records only purchases and payments they made.')
body(doc,'Verification: the local static preview loaded the sign-in interface without console errors. Public deployment is intentionally pending: Supabase Authentication must first be configured with the GitHub Pages URL as its Site URL and redirect URL, then the implementation needs a signed-in two-device verification.')

heading(doc,'13 July 2026, 14:28 IST — Supabase function-permission hardening')
body(doc,'Finding: Supabase Security Advisor correctly reported that PostgreSQL functions can receive broad EXECUTE privileges by default. The initial schema used SECURITY DEFINER for controlled household creation/join and its RLS membership helper, so the default grants produced warnings even though each function checked auth.uid().')
body(doc,'What changed: added supabase/harden-function-permissions.sql and updated the baseline schema. The hardening migration revokes EXECUTE from public, anonymous, and ordinary signed-in roles for all existing public functions (including the dashboard-generated rls_auto_enable helper), prevents future public functions from inheriting broad permissions, moves the RLS helper to a non-exposed private schema, and re-grants only authenticated access to the guarded create_household and join_household RPCs.')
body(doc,'Required operator action: run the hardening migration once in Supabase SQL Editor, then recheck Security Advisor. The advisor may still flag the two signed-in SECURITY DEFINER RPCs; that is expected because signed-in users must call them. Their scope is restricted by auth.uid(), fixed search paths, RLS, and explicit grants. Two-device authentication and data-sync verification remain pending.')

heading(doc,'13 July 2026, 14:50 IST — Cross-browser email-code sign-in')
body(doc,'Finding: magic links establish a session in the browser that consumes the link. That is normal passwordless-auth behaviour but fails the project requirement that either person should be able to choose any browser/device without a browser handoff.')
body(doc,'What changed: replaced the web client’s magic-link-only screen with an email one-time-code flow. The user requests a code, then enters the code in the chosen browser; the client verifies it with Supabase Auth. Added supabase/email-otp-template.html for the Supabase Magic Link email template, which sends the code and a neutral link to open the website rather than a session-carrying link.')
body(doc,'Required operator action and verification limit: paste the supplied template into Supabase Authentication > Email Templates > Magic Link and save it, then send a new code on each Mac. The client code passes JavaScript syntax validation; live two-device sign-in, household join, purchase sync, and settlement sync must be verified after deployment.')

heading(doc,'13 July 2026, 15:15 IST — Free-plan email-template correction')
body(doc,'Finding: the Supabase Free project’s default email sender does not permit custom email-template edits; the dashboard correctly requires custom SMTP for that. Therefore the proposed code-template route cannot be used without introducing another provider or cost.')
body(doc,'Correction: restored the supported default magic-link flow and made the browser boundary explicit in the sign-in UI. Any browser is supported: copy the unconsumed email sign-in link and paste it into the address bar of the browser where the session is wanted. The link then creates the session there. No custom SMTP, password, or additional data collection is required.')
body(doc,'Verification status: this restores the official Free-plan-compatible flow. The corrected JavaScript syntax must be checked and the public website redeployed before the two-Mac authentication and shared-ledger test is repeated.')

heading(doc,'13 July 2026, 15:35 IST — Shared-login feedback and database-access repair')
body(doc,'Finding: a live test showed no visible confirmation after requesting a sign-in link, leaving the action ambiguous. A separate signed-in second Mac displayed “permission denied for table household_members.” The latter is a table-privilege gap: RLS policies existed, but authenticated Data API requests also require table-level privileges before those policies can evaluate.')
body(doc,'What changed: the sign-in panel now displays immediate in-place Sending, Sent, and error states and disables the button while the request is in flight. Added explicit authenticated SELECT access to the four reviewed-ledger tables and only the write grants used by purchases/settlements. RLS remains enabled and continues to limit rows to household members. Added supabase/grant-ledger-table-access.sql for the already-created database and updated the baseline schema.')
body(doc,'Required operator action and verification limit: run the new grant repair once in the Supabase SQL Editor, then deploy the web update. JavaScript syntax is checked locally; actual two-Mac membership loading, invite joining, expense creation, settlement creation, and real-time sync must be verified after the SQL and deployment complete.')

heading(doc,'13 July 2026, 15:50 IST — Sign-in UI clarity pass')
body(doc,'What changed: shortened the sign-in explanation, added a focused responsive sign-in layout, and gave the feedback state a clear visual hierarchy. The button now changes to Sending while disabled; success is a green message card; errors are an orange message card. The Supabase Free email-rate error is translated into an actionable wait/use-your-newest-link explanation instead of raw API wording.')
heading(doc,'13 July 2026, 16:05 IST — Sign-in retry time')
body(doc,'What changed: replaced the generic email-rate-limit explanation with a disabled retry button and a specific local time for the next request. The time is calculated from the one-hour Free-plan window after the browser receives a rate-limit response, is remembered across reloads in that browser, and re-enables automatically. It is explicitly a client-side estimate because the email service does not provide a reset timestamp.')
heading(doc,'13 July 2026, 16:15 IST — Data API privilege verification repair')
body(doc,'Correction: the earlier table-grant migration returned SQL success but a signed-in second Mac still received “permission denied for table household_members.” Therefore SQL completion alone is not accepted as verification. The repair now also grants authenticated usage on the public schema, asks PostgREST to reload its schema/privilege cache, and returns four effective-privilege checks. The shared-ledger test remains incomplete until those checks are true and the second Mac can read its membership without an error.')
heading(doc,'13 July 2026, 16:25 IST — Second-Mac membership read verified')
body(doc,'Verification: the effective-privilege query returned true for public-schema usage, membership read, purchase access, and settlement access. After refresh, the second Mac reached the signed-in “Create or join a household” state with no table-permission message. This confirms the membership-read repair; household creation, invite joining, expense creation, and cross-device refresh remain the next live checks.')
heading(doc,'13 July 2026, 16:35 IST — Owner email invite action')
body(doc,'What changed: a household owner now sees Email invite and Copy code actions after creating or joining their household. Email invite opens the device’s own mail composer with the project URL and invite code prefilled; it does not send mail automatically and Grocery Ledger stores no recipient email address or email content. This lets the Mac 2 owner invite Mac 1 without using the rate-limited authentication email service.')
body(doc,'Why: live screenshots showed that the original message was easy to miss and the unchanged button made a successful send feel unresponsive. This pass makes the authentication state visible at the point of action without changing the data, privacy, or permission model.')

heading(doc,'13 July 2026, 16:50 IST — Local-only PDF import with atomic duplicate protection')
body(doc,'What changed: the shared website now offers Import PDF after a person has joined a household. PDF.js reads a selected receipt entirely in that browser; the user must review and edit the proposed label, amount, date, category, personal flag, and restock flag before saving. The raw PDF and its extracted text are discarded, never uploaded to Supabase, and never put in the GitHub repository.')
body(doc,'Duplicate safeguard: the browser derives two SHA-256 fingerprints locally—one from the PDF bytes and one from normalized extracted text. Supabase stores only those fingerprints in invoice_imports. A guarded import_purchase database function atomically reserves both fingerprints and creates the reviewed expense, so two members importing the same bill at the same time cannot double-count it. It also catches a re-exported/copy PDF when its readable receipt content is unchanged.')
body(doc,'Limit and operator action: the first PDF pass can only suggest a date and likely total, so the review step is mandatory and it does not claim line-item accuracy. The owner must run supabase/add-local-pdf-imports.sql once before this feature can save imports. JavaScript syntax and local browser loading are verified; a full end-to-end upload must be performed after the migration on the live site.')

heading(doc,'14 July 2026, 13:14 IST — Multi-household lifecycle and permission model agreed')
body(doc,'Decision: one signed-in account may belong to multiple households. A visible household name is not unique and must never be used as identity; the existing UUID household ID and unique invite code identify the household, so different people may safely use the same name.')
body(doc,'Approved roles: Owner manages everything; Owner or explicitly appointed Admin may add/remove members; Members may add, edit, or delete only entries they created and may request Admin access from the Owner. Owner alone controls role changes, archive, restore, and permanent deletion. A member with an unpaid balance must settle before removal so balances cannot become orphaned.')
body(doc,'Lifecycle: archive is a read-only, recoverable household state; restore returns it to normal use. Permanent deletion follows a recovery period and must purge the household ledger rather than auth identities. The current test household has not been removed while the lifecycle design is under discussion.')
body(doc,'Flow decision pending implementation: sign-in is a dedicated first screen. After authentication, users without a household see only Create or Join; users with households land in the last active household and use a household switcher. Dashboard, expenses, restock, settlements, and settings do not render behind authentication or onboarding.')
body(doc,'Reference: current Splitwise help describes group-wide bill editing/deletion and no special admin roles, plus group removal/deletion/recovery flows. Grocery Ledger intentionally differs with Owner/Admin/Member controls and an activity audit trail, while retaining the clear group-settings and settled-before-removal concepts.')

heading(doc,'14 July 2026, 14:05 IST — Approved multi-household lifecycle implemented for web handoff')
body(doc,'What changed: added supabase/multi-household-lifecycle.sql as a new, idempotent migration for the existing project. It adds Owner/Admin/Member roles, multiple household membership per account, member admin requests, owner-only approval and ownership transfer, manager-controlled member removal after the member balance is settled, and a narrow activity table containing action names and opaque IDs only.')
body(doc,'Lifecycle: a household may be archived only by its Owner when every active member balance is zero. Archive is read-only and sets a 30-day recovery deadline. Only the Owner can restore within that period; permanent deletion is deliberately unavailable until the deadline and purges the household ledger/memberships rather than an auth identity. Visible household names are explicitly non-unique.')
body(doc,'Website flow: replaced the mixed sign-in/dashboard rendering with separate sign-in, household-picker, and selected-household dashboard states. A signed-in account can create, join, select, and switch among multiple households. The dashboard exposes local PDF import, entry archive/restore, household settings, invite actions, role requests, and archive/recovery controls according to role.')
body(doc,'Verification and operator action: JavaScript syntax checks passed with the bundled Node runtime; four automated static tests passed, covering staged flow, role/recovery SQL, privacy PDF safeguards, and the prior web controls. The new migration has not yet run against the live Supabase project, so it must be run once in a new SQL Editor tab before the live Page can exercise these controls. Do not rerun schema.sql because its original policy names are intentionally one-time only.')

heading(doc,'14 July 2026, 14:25 IST — Lifecycle migration retry hardening')
body(doc,'Finding: applying the lifecycle migration to the existing Supabase project stopped at a permission-cleanup statement because an older remove_household_member function signature was absent. The database therefore did not complete the lifecycle migration. This was an avoidable migration-compatibility defect, not a household deletion.')
body(doc,'Repair: permission cleanup now checks whether each legacy function signature exists before attempting a revoke, so both early test databases and the current database can proceed safely. The migration also now grants and guards settlement updates, which are required for the promised archive/restore behaviour of settlement entries.')
body(doc,'Verification and next action: JavaScript syntax, SQL whitespace checks, and four automated static tests pass locally. This change has not yet been applied to Supabase or pushed to GitHub Pages. Re-open the local multi-household-lifecycle.sql file, replace the failed SQL Editor tab contents with the corrected version, and run it once; do not run the baseline schema again.')

heading(doc,'14 July 2026, 14:38 IST — Lifecycle grant signature correction')
body(doc,'Finding: the retry run exposed a second signature mismatch in the same migration. The final authenticated EXECUTE grant named remove_household_member(uuid), while the function correctly takes both a household UUID and a member UUID. PostgreSQL therefore stopped the migration at the grant statement.')
body(doc,'Repair: corrected the grant to remove_household_member(uuid, uuid). The automated lifecycle test now explicitly asserts that the exact two-argument grant is present, in addition to the legacy-safe conditional revoke.')
body(doc,'Verification and next action: this correction is verified only by static checks until the complete revised file is run in Supabase. Replace the whole SQL Editor contents from the local migration file again; do not try to run only the visible final lines, because the migration is designed to reapply safely as a complete unit.')
doc.save(OUT)
print(OUT)
