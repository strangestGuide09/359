from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Grocery_Slip_Tracker_Market_Research.docx'

def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tcPr = cell._tc.get_or_add_tcPr(); node = tcPr.first_child_found_in('w:tcMar')
    if node is None:
        node = OxmlElement('w:tcMar'); tcPr.append(node)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = node.find(qn('w:' + side))
        if el is None:
            el = OxmlElement('w:' + side); node.append(el)
        el.set(qn('w:w'), str(value)); el.set(qn('w:type'), 'dxa')

def set_cell_text(cell, text, bold=False, color=None, size=9.2):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text); set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14 if level == 1 else 9); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_font(r, 16 if level == 1 else 12.5, True, (46,116,181) if level == 1 else (31,77,120))
    return p

def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.1
    if bold_lead:
        r = p.add_run(bold_lead); set_font(r, 10.5, True)
    r = p.add_run(text); set_font(r, 10.5)
    return p

def add_bullet(doc, text, size=10.2):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(text), size)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.width = Inches(widths[i]); shade(cell, 'EAF1F8'); set_cell_text(cell, h, True, (31,77,120), 9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i]); set_cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

doc = Document(); sec = doc.sections[0]
sec.top_margin = Inches(.8); sec.bottom_margin = Inches(.75); sec.left_margin = Inches(.8); sec.right_margin = Inches(.8)
doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); doc.styles['Normal'].font.size = Pt(10.5)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run('PROJECT HUB  |  PRODUCT & DESIGN REFERENCE'), 8.5, True, (89,89,89))
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run('Grocery Slip & Shared Household Tracker  •  Decision reference'), 8, False, (89,89,89))

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('Grocery Slip & Shared Household Tracker'), 23, True, (11,37,69))
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
set_font(p.add_run('Product research, requirements, architecture, and critical decisions'), 13, False, (85,85,85))

add_table(doc, ['Document role', 'Rule'], [
    ['This document', 'A durable reference for research, product scope, requirements, flows, data design, privacy, and approved decisions. It is updated when those decisions change.'],
    ['GroceryLedger Change Log', 'The dated implementation record: code changes, migrations, tests, releases, verification outcomes, defects, and operational follow-up.'],
], [1.75, 4.95])

add_heading(doc, 'Executive view', 1)
add_body(doc, 'The idea has genuine value for roommates, couples, and families who repeatedly buy mixed groceries: one receipt should create a fair shared balance and a useful “possible buys” list. The market is not empty—Tricount covers free group splitting, Splitwise covers receipt/item splitting as a paid feature, and Grocy/pantry apps cover stock, barcodes, expiry, and shopping lists. The opportunity is the integrated, privacy-conscious bridge from reviewed receipt to shared ledger to low-confidence replenishment cue.')
add_body(doc, 'Do not market this as an all-knowing pantry AI. A receipt proves a purchase, not consumption or remaining stock. The product should be explicit about evidence, ask for lightweight confirmation, and describe outputs as possible buys or estimates.', 'Honest product constraint: ')

add_heading(doc, 'Confirmed product scope', 1)
add_table(doc, ['Area', 'Approved direction'], [
    ['Initial users', 'Ekta and Ritesh in India. INR is the initial currency.'],
    ['Surfaces', 'iPhone app and a browser version for multi-device shared use.'],
    ['Primary jobs', 'Split shared groceries and household costs; track balances and settlements; review likely restocks.'],
    ['Expense types', 'Groceries, Food, Wi-Fi, Water, Household, and Other. Food/bills affect balances but do not create grocery-restock evidence.'],
    ['Receipt input', 'PDF first. PDFs are processed locally in the chosen browser/device; raw PDFs and extracted receipt text are not stored or synced.'],
    ['Shared-item rule', 'Equal split is the default for a shared expense. A reviewed item/expense can be personal and excluded from the other member’s share.'],
    ['Settlements', 'Record a Splitwise-style payment between members; do not collect money or retain payment instruments/references.'],
    ['Prediction', 'Suggestions, not certainty. Use the latest two distinct purchase dates for a tracked item; later incorporate more confirmed evidence.'],
    ['Notifications', 'User chooses frequency. Message is neutral (“review possible buys”), never a claim that an item is definitely exhausted.'],
], [1.45, 5.25])

add_heading(doc, 'Privacy and data-minimisation policy', 1)
add_body(doc, 'The system stores reviewed ledger information only. It intentionally does not store raw PDFs, raw extracted receipt text, street address, payment method, payment reference, payment status, bank/card/UPI information, or a recipient email for household invites. Authentication identity is handled by Supabase Auth; an email address is not copied into the ledger tables.')
add_table(doc, ['May be retained', 'Must not be retained'], [
    ['Household UUID/name; membership role; reviewed merchant/label; category; amount; date; payer identity; personal/restock flags; estimate date; settlement amount/date; opaque duplicate fingerprints; audit action/type/time.', 'Receipt files or receipt text; customer/address information; payment method or transaction reference; bank/card/UPI information; payment status; email invite recipient or email body.'],
], [3.35, 3.35])
add_body(doc, 'Local PDF duplicate protection uses hashes of the selected file and normalized local receipt text. The database keeps only the two hashes and the reviewed purchase; this prevents a second member importing the same bill from double-counting it without uploading the source receipt.', 'Duplicate rule: ')

add_heading(doc, 'Market research and positioning', 1)
add_table(doc, ['Product', 'What it already solves', 'Gap Grocery Ledger should address'], [
    ['Tricount', 'Free shared expenses, custom splits, settle-up.', 'No receipt-to-inventory/replenishment workflow.'],
    ['Splitwise', 'Group debts and paid receipt itemisation.', 'No household restock intelligence; broad group-edit model is unsuitable for the selected permission model.'],
    ['Grocy', 'Open-source inventory, stock thresholds, barcodes, expiry.', 'High entry effort and no receipt-led shared-balance focus.'],
    ['Pantry Check / NoWaste', 'Inventory, expiry and shopping list cues.', 'No integrated roommate settlement or personal-item allocation.'],
], [1.25, 2.4, 3.05])
add_body(doc, 'Position the product as a “shared household ledger with gentle restock cues.” The differentiator is not receipt parsing alone. It is a transparent, low-effort loop: receipt/manual entry → review and ownership correction → balance → repeat-item evidence → suggested next buy.', 'Recommended thesis: ')

add_heading(doc, 'Experience flow', 1)
add_body(doc, 'The site deliberately uses stages so that authentication and household management do not compete with the ledger.')
add_table(doc, ['Stage', 'What the user sees', 'Primary action'], [
    ['1. Account gate', 'Only sign up/sign in and the privacy promise.', 'Request a passwordless email link and open it in the intended browser.'],
    ['2. Household picker', 'Only the account’s active households plus Create / Join.', 'Create a new household or join with an invite code.'],
    ['3. Selected ledger', 'Balances, expenses, settlements, restock cues, and an active-household switcher.', 'Add/review an expense, import a local PDF, settle, or switch household.'],
    ['4. Household settings', 'Members, roles, invite actions, admin requests, archive/recovery controls.', 'Invite, copy code, manage membership or lifecycle according to role.'],
], [1.25, 3.3, 2.15])
add_body(doc, 'A successful browser sign-in persists through reloads, tab switches, and normal browser reopening. Sign-in is required again only after explicit sign-out, cleared site data, or an expired/revoked session. This is safer and less frustrating than forcing a sign-in whenever a tab closes.', 'Session decision: ')

add_heading(doc, 'Household roles and lifecycle', 1)
add_table(doc, ['Role', 'Permissions'], [
    ['Owner', 'Manages all household settings, members, role decisions, archive/restore, recovery, ownership transfer, and permanent deletion after the recovery period.'],
    ['Admin', 'May invite/add/remove members and manage household membership. An Admin is appointed by Owner approval.'],
    ['Member', 'May add, edit, archive, restore, or delete only their own expenses/settlements; may request Admin access.'],
], [1.2, 5.5])
add_bullet(doc, 'Visible household names are not unique. Household UUIDs and invite codes are the identifiers, so different users may safely use the same name.')
add_bullet(doc, 'One account may belong to multiple households and can switch between them.')
add_bullet(doc, 'A member cannot be removed while their balance is unsettled; this prevents orphaned debt.')
add_bullet(doc, 'Archive makes a household read-only. The Owner can restore it during a 30-day recovery period; permanent deletion is allowed only afterwards and never deletes an authentication identity.')

add_heading(doc, 'ERD — approved logical model', 1)
add_body(doc, 'This is the logical data model. It records only reviewed ledger data and supporting access/lifecycle records; it is not a receipt-storage model.')
add_table(doc, ['Entity', 'Key relationships'], [
    ['Auth user', 'Supabase Auth identity. One user has many household memberships; no duplicate profile/email table is required for the ledger.'],
    ['Household', 'One household has many members, invites, purchases, settlements, import fingerprints, and activity records.'],
    ['Household member', 'Joins one user to one household and assigns Owner, Admin, or Member role.'],
    ['Purchase', 'Belongs to one household and is created by one member. Carries reviewed split/restock/category fields; may reference one local-only import reservation.'],
    ['Settlement', 'Belongs to one household; records payer, receiver, amount and date.'],
    ['Household invite', 'Belongs to one household; created by Owner/Admin; code-based join without storing an invitee email.'],
    ['Admin request / activity', 'Belongs to one household; captures a role request or minimal action trail using opaque IDs, not receipt content.'],
    ['Invoice import', 'Belongs to one household and keeps file/content fingerprints only; enforces one reviewed import per bill.'],
], [1.65, 5.05])

add_heading(doc, 'DRD — essential record definitions', 1)
add_table(doc, ['Record', 'Required attributes / purpose'], [
    ['Household', 'id (UUID), name (non-unique), created_by, created_at, archive/recovery fields.'],
    ['Membership', 'household_id, user_id, role, joined_at. Unique by household + user.'],
    ['Purchase', 'id, household_id, created_by, label/merchant, category, amount, purchase date, paid_by, is_personal, track_for_restock, optional estimated_use_by, created/archived metadata.'],
    ['Settlement', 'id, household_id, created_by, payer, receiver, amount, settled date, archived metadata. No payment mode/reference field.'],
    ['Invite', 'id, household_id, code, created_by, created/expires/used metadata. No recipient email field.'],
    ['Import fingerprint', 'household_id, file_hash, content_hash, purchase_id. Hashes support duplicate protection only; no file/text payload.'],
], [1.55, 5.15])

add_heading(doc, 'Restock and expiry design', 1)
add_bullet(doc, 'A grocery/household item must be explicitly tracked for restock; Food and personal items are excluded.')
add_bullet(doc, 'The first cadence estimate is the interval between the latest two distinct purchase dates. Same-day duplicate evidence withholds a prediction.')
add_bullet(doc, 'With longer confirmed history, present a robust median/range rather than a simplistic average, together with the evidence count.')
add_bullet(doc, 'A user-entered estimated use-by date takes precedence over cadence. It is always labelled “estimated,” not a known package expiry.')
add_bullet(doc, 'Later enrichment may use barcode/product catalog information or package-label scans, but it must be opt-in and must distinguish generic shelf-life guidance from the item’s actual expiry date.')

add_heading(doc, 'Open design questions', 1)
add_table(doc, ['Question', 'Current direction'], [
    ['How should invite delivery work?', 'Use copy-code and a prefilled device mail draft; the app does not send/store recipient email.'],
    ['How should barcode lookup be introduced?', 'Only after a consented product-data provider is selected, with a source/last-updated label and no claim that it provides the actual package expiry.'],
    ['What should notification defaults be?', 'No default frequency. Each user chooses Off, Daily, or Weekly and a time.'],
    ['When should item-level parsing be trusted?', 'Never silently. Parsing remains a reviewed draft; low-confidence labels/totals need explicit confirmation.'],
], [2.55, 4.15])

add_heading(doc, 'Sources', 1)
for name, url, note in [
    ('Tricount', 'https://tricount.com/en-in/ | https://tricount.com/expense-tracker-features', 'Shared-expense positioning and split features.'),
    ('Splitwise', 'https://www.splitwise.com/pro | https://www.splitwise.com/subscriptions/new', 'Receipt scanning/itemisation and subscription positioning.'),
    ('Grocy', 'https://grocy.info/', 'Open-source stock, barcode, expiry, and shopping-list functionality.'),
    ('Pantry Check / NoWaste', 'https://pantrycheck.com/ | https://www.nowasteapp.com/', 'Consumer pantry and expiry/restock approaches.'),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(name + ' — '), 8.5, True, (31,77,120)); set_font(p.add_run(note + ' ' + url), 8.5, False, (70,70,70))

doc.save(OUT)
print(OUT)
